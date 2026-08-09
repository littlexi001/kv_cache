from __future__ import annotations

import copy
import importlib.util
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = (
    PROJECT_ROOT
    / "doc"
    / "nine_newline_soft_threshold_analysis_20260727.docx"
)
FIGURE_PATH = (
    PROJECT_ROOT
    / "figures"
    / "nine_newline_probability_136k_144k_20260727.png"
)
TABLE_GEOMETRY_PATH = Path(
    r"C:\Users\27814\.codex\plugins\cache\openai-primary-runtime"
    r"\documents\26.723.12215\skills\documents\scripts\table_geometry.py"
)


def load_table_geometry():
    spec = importlib.util.spec_from_file_location("table_geometry", TABLE_GEOMETRY_PATH)
    if spec is None or spec.loader is None:
        raise RuntimeError("could not load table_geometry.py")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


table_geometry = load_table_geometry()


INK = "172033"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
LIGHT_BLUE = "EEF5FB"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D0D5DD"
WHITE = "FFFFFF"
ORANGE = "F28C38"
GREEN = "13795B"
RED = "B42318"
BODY_FONT = "Calibri"
CJK_FONT = "Microsoft YaHei"
MATH_FONT = "Cambria Math"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(
    run,
    *,
    name: str = BODY_FONT,
    east_asia: str = CJK_FONT,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = MID_GRAY, size: int = 4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_paragraph_fill_and_left_border(
    paragraph,
    *,
    fill: str = LIGHT_BLUE,
    border_color: str = BLUE,
) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border_color)
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), "140")
    spacing.set(qn("w:after"), "140")


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))


def add_custom_bullet_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    existing_abstract = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(existing_abstract, default=-1) + 1
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract_num.append(multi_level)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level_justification = OxmlElement("w:lvlJc")
    level_justification.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "720")
    indentation.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend((tabs, indentation, spacing))
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Symbol")
    fonts.set(qn("w:hAnsi"), "Symbol")
    r_pr.append(fonts)
    level.extend((start, num_fmt, level_text, level_justification, p_pr, r_pr))
    abstract_num.append(level)
    numbering.append(abstract_num)

    existing_nums = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    num_id = max(existing_nums, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def set_list_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num_id_element))


def add_bullet(doc: Document, num_id: int, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_list_numbering(paragraph, num_id)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=INK)


def add_body(
    doc: Document,
    text: str,
    *,
    bold_prefix: str | None = None,
    after: float = 6,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run_font(prefix, size=11, color=INK, bold=True)
        body = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(body, size=11, color=INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=INK)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    # The paragraph is initially empty in python-docx except for this run.
    set_run_font(
        run,
        size={1: 16, 2: 13, 3: 12}[level],
        color=BLUE if level < 3 else DARK_BLUE,
        bold=True,
    )


def add_callout(doc: Document, label: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.10)
    paragraph.paragraph_format.line_spacing = 1.10
    set_paragraph_fill_and_left_border(paragraph)
    label_run = paragraph.add_run(f"{label}  ")
    set_run_font(label_run, size=10.5, color=DARK_BLUE, bold=True)
    body_run = paragraph.add_run(text)
    set_run_font(body_run, size=10.5, color=INK)


def math_run(text: str, *, plain: bool = False):
    run = OxmlElement("m:r")
    r_pr = OxmlElement("m:rPr")
    if plain:
        style = OxmlElement("m:sty")
        style.set(qn("m:val"), "p")
        r_pr.append(style)
    run.append(r_pr)
    text_element = OxmlElement("m:t")
    text_element.text = text
    run.append(text_element)
    return run


def math_sub(base, subscript):
    element = OxmlElement("m:sSub")
    properties = OxmlElement("m:sSubPr")
    base_container = OxmlElement("m:e")
    sub_container = OxmlElement("m:sub")
    append_math(base_container, base)
    append_math(sub_container, subscript)
    element.extend((properties, base_container, sub_container))
    return element


def math_sup(base, superscript):
    element = OxmlElement("m:sSup")
    properties = OxmlElement("m:sSupPr")
    base_container = OxmlElement("m:e")
    sup_container = OxmlElement("m:sup")
    append_math(base_container, base)
    append_math(sup_container, superscript)
    element.extend((properties, base_container, sup_container))
    return element


def math_fraction(numerator, denominator):
    element = OxmlElement("m:f")
    properties = OxmlElement("m:fPr")
    numerator_container = OxmlElement("m:num")
    denominator_container = OxmlElement("m:den")
    append_math(numerator_container, numerator)
    append_math(denominator_container, denominator)
    element.extend((properties, numerator_container, denominator_container))
    return element


def append_math(container, value) -> None:
    if isinstance(value, (list, tuple)):
        for item in value:
            append_math(container, item)
    else:
        container.append(value)


def add_equation(doc: Document, elements, *, after: float = 7) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.keep_together = True
    equation = OxmlElement("m:oMath")
    append_math(equation, elements)
    paragraph._p.append(equation)


def sub(symbol: str, script: str):
    return math_sub(math_run(symbol), math_run(script, plain=True))


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
    *,
    numeric_columns: set[int] | None = None,
    font_size: float = 9.5,
) -> None:
    numeric_columns = numeric_columns or set()
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "1")
    table.rows[0]._tr.get_or_add_trPr().append(table_header)
    for column, header in enumerate(headers):
        cell = table.rows[0].cells[column]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
            if column in numeric_columns
            else WD_ALIGN_PARAGRAPH.LEFT
        )
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=9.5, color=INK, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for column, value in enumerate(values):
            cell = cells[column]
            set_cell_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if column in numeric_columns
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            run = paragraph.add_run(value)
            color = RED if value.startswith("61.13") or value.startswith("99.6") else INK
            set_run_font(run, size=font_size, color=color)
    table_geometry.apply_table_geometry(
        table,
        widths_dxa,
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 100, "bottom": 100, "start": 120, "end": 120},
    )
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        1: (16, BLUE, 16, 8),
        2: (13, BLUE, 12, 6),
        3: (12, DARK_BLUE, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.paragraph_format.space_after = Pt(0)
    tabs = header_paragraph.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.5), alignment=2)
    left = header_paragraph.add_run("LONG-CONTEXT RETRIEVAL")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = header_paragraph.add_run("\tQwen3-8B · RESEARCH NOTE")
    set_run_font(right, size=8.5, color=MUTED)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_paragraph.paragraph_format.space_before = Pt(0)
    label = footer_paragraph.add_run("Page ")
    set_run_font(label, size=8.5, color=MUTED)
    add_page_number(footer_paragraph)


def add_title_block(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(5)
    kicker_run = kicker.add_run("RESEARCH NOTE  ·  LONG-CONTEXT RETRIEVAL")
    set_run_font(kicker_run, size=9, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    title_run = title.add_run("从 nine 到 newline：证据主导权如何被夺走")
    set_run_font(title_run, size=23, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle_run = subtitle.add_run(
        "136K–144K 逐 token 扫描、关键 QK 软阈值与输出硬边界"
    )
    set_run_font(subtitle_run, size=12.5, color=MUTED)

    metadata = doc.add_paragraph()
    metadata.paragraph_format.space_before = Pt(0)
    metadata.paragraph_format.space_after = Pt(13)
    metadata_run = metadata.add_run(
        "模型：Qwen3-8B    样本：8193 个连续长度点    日期：2026-07-27"
    )
    set_run_font(metadata_run, size=9.5, color=MUTED)


def build_document() -> Document:
    if not FIGURE_PATH.exists():
        raise FileNotFoundError(FIGURE_PATH)
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    bullet_num_id = add_custom_bullet_numbering(doc)
    add_title_block(doc)

    add_callout(
        doc,
        "核心结论",
        "输出层存在严格的零阈值；模型内部则存在一个很强、但并非绝对的经验软阈值。"
        "这里将此前口头表述的 newnine 统一记为强竞争输出 token：newline。",
    )

    figure_paragraph = doc.add_paragraph()
    figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_paragraph.paragraph_format.space_before = Pt(8)
    figure_paragraph.paragraph_format.space_after = Pt(2)
    figure_paragraph.paragraph_format.keep_with_next = True
    run = figure_paragraph.add_run()
    figure = run.add_picture(str(FIGURE_PATH), width=Inches(6.35))
    figure._inline.docPr.set(
        "descr",
        "P(nine) 与 P(newline) 从 136K 到 144K 的逐 token 概率曲线；"
        "标注首次平局、首次连续五点失败和持续失效窗口。",
    )
    figure._inline.docPr.set("title", "nine 与 newline 概率变化")
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(8)
    caption_run = caption.add_run(
        "图 1  P(nine) 与 P(newline) 的逐 token 变化。虚线标出首次平局、"
        "首次连续 5 点失败和持续失效窗口。"
    )
    set_run_font(caption_run, size=9, color=MUTED, italic=True)

    add_heading(doc, "1. 输出层存在严格的决策阈值", 1)
    add_body(
        doc,
        "令 z 表示最终 LM Head 给输出 token 的 logit。nine 相对 newline 的输出 margin 为：",
    )
    add_equation(
        doc,
        [
            sub("Δ", "out"),
            math_run(" = "),
            sub("z", "nine"),
            math_run(" − "),
            sub("z", "newline"),
            math_run(" = log "),
            math_fraction(
                [math_run("P", plain=True), math_run("(nine)", plain=True)],
                [math_run("P", plain=True), math_run("(newline)", plain=True)],
            ),
        ],
    )
    add_bullet(doc, bullet_num_id, "Δout > 0：nine 的 logit 更高，模型输出 nine。")
    add_bullet(doc, bullet_num_id, "Δout < 0：newline 的 logit 更高，模型输出 newline。")
    add_bullet(doc, bullet_num_id, "Δout = 0：两者恰好位于决策边界。")
    add_body(
        doc,
        "因此，模型真正执行的硬开关是输出 margin 是否穿过零点。"
        "这个结论是 softmax 定义直接给出的，不依赖经验拟合。",
    )

    add_heading(doc, "2. 主导权并非一次性、永久地交接", 1)
    add_body(
        doc,
        "完整扫描覆盖 136K–144K，总计 8193 个连续长度点。"
        "结果表现为大量恢复与再次失效，而不是随长度单调下降。",
    )
    add_table(
        doc,
        ["边界定义", "判定标准", "新增 token", "对应总长度"],
        [
            ["瞬时边界", "首次 Δout ≤ 0", "90", "136.088K"],
            ["短段稳定失败", "连续 5 点 Δout ≤ 0", "1886", "137.842K"],
            ["统计边界", "64 点窗口失败率 ≥ 50%", "4160", "140.063K"],
            ["持续失效窗口", "256 点窗口失败率 ≥ 80%", "4823", "140.710K"],
        ],
        [1900, 3560, 1800, 2100],
        numeric_columns={2, 3},
    )
    add_body(
        doc,
        "全区间共发生 907 次正确/错误翻转，newline 获胜比例为 25.22%。"
        "因此，更准确的描述是“带有大量抖动的相变区域”，而不是某个固定长度之后永久失败。",
    )

    add_heading(doc, "3. 主要是 nine 的支持崩塌，而非 newline 暴涨", 1)
    add_table(
        doc,
        ["区间", "平均 P(nine)", "平均 P(newline)", "newline 获胜率", "关键 QK"],
        [
            ["136–140K", "46.74%", "14.75%", "6.54%", "15.11"],
            ["140–141K", "20.76%", "20.09%", "61.13%", "13.20"],
            ["141–144K", "28.21%", "17.57%", "38.12%", "13.83"],
        ],
        [1500, 1900, 2050, 2050, 1860],
        numeric_columns={1, 2, 3, 4},
    )
    add_body(
        doc,
        "进入主要失败区时，P(nine) 平均下降约 26 个百分点，而 P(newline) 只上升约 5 个百分点。"
        "这说明主导权转移的主要来源不是 newline 概率突然暴涨，而是模型从真实证据得到的 nine 支持显著减弱；"
        "newline 随后依靠原本存在的格式结束/换行先验接管输出。",
    )

    add_heading(doc, "4. 最强内部预警指标：关键 head 的加权 QK", 1)
    add_body(
        doc,
        "对此前识别出的 29 个关键 head，使用其正确上下文证据 mass 作为权重，定义：",
    )
    add_equation(
        doc,
        [
            sub("S", "QK"),
            math_run(" = "),
            math_sub(math_run("∑"), math_run("h∈ℋ", plain=True)),
            math_run(" "),
            sub("w", "h"),
            math_run(" "),
            math_fraction(
                [
                    math_sup(sub("q", "h"), math_run("T", plain=True)),
                    math_run(" "),
                    sub("k", "nine,h"),
                ],
                [math_run("√d")],
            ),
        ],
    )
    add_equation(
        doc,
        [sub("S", "QK"), math_run(" ≈ 14.11")],
    )
    add_body(
        doc,
        "14.11 是从本次数据中得到的经验软阈值：它能很好地预警失败，但并不是网络中实际存在的硬开关。",
    )
    add_bullet(doc, bullet_num_id, "SQK ≤ 14.11 时，newline 获胜概率为 67.2%。")
    add_bullet(doc, bullet_num_id, "SQK > 14.11 时，newline 获胜概率只有 4.8%。")
    add_bullet(doc, bullet_num_id, "Balanced accuracy 为 86.4%；与输出 margin 的相关系数 r = 0.839。")
    add_bullet(doc, bullet_num_id, "单独解释输出 margin 方差的 R² 为 70.4%。")
    add_table(
        doc,
        ["加权 QK 区间", "newline 获胜概率", "解释"],
        [
            ["≥ 15", "0.7%", "真实证据占据明显优势"],
            ["14.11–15", "10.3%", "大多数点仍由 nine 主导"],
            ["13–14.11", "46.3%", "进入不稳定相变区域"],
            ["12–13", "90.8%", "newline 几乎稳定接管"],
            ["< 12", "99.6%", "真实证据支持基本失效"],
        ],
        [2300, 2300, 4760],
        numeric_columns={0, 1},
    )
    add_callout(
        doc,
        "如何理解 14.11",
        "它是“失败风险急剧增加”的预警线，而不是“低于后必然失败”的因果硬阈值。"
        "在 14.11 以下仍有恢复点，在 14.11 以上也存在少量瞬时失败。",
    )

    add_heading(doc, "5. 从 QK 下降到输出切换：完整计算链条", 1)
    add_heading(doc, "5.1 QK 分数控制证据的相对读取优势", 2)
    add_body(doc, "在某个 attention head 中，真实证据 token g 的 attention 为：")
    add_equation(
        doc,
        [
            sub("a", "g"),
            math_run(" = "),
            math_fraction(
                [math_run("exp", plain=True), math_run("("), sub("s", "g"), math_run(")")],
                [
                    math_run("exp", plain=True),
                    math_run("("),
                    sub("s", "g"),
                    math_run(") + "),
                    math_sub(math_run("∑"), math_run("j≠g", plain=True)),
                    math_run(" exp", plain=True),
                    math_run("("),
                    sub("s", "j"),
                    math_run(")"),
                ],
            ),
        ],
    )
    add_body(doc, "真实证据相对某个竞争 token c 的 attention 比例为：")
    add_equation(
        doc,
        [
            math_fraction(sub("a", "g"), sub("a", "c")),
            math_run(" = exp", plain=True),
            math_run("("),
            sub("s", "g"),
            math_run(" − "),
            sub("s", "c"),
            math_run(")"),
        ],
    )
    add_body(
        doc,
        "如果真实证据相对竞争 token 的 QK 分数差降低 1，而其他分数不变，则：",
    )
    add_equation(
        doc,
        [
            math_fraction(
                math_sup(sub("a", "g"), math_run("′")),
                math_sup(sub("a", "c"), math_run("′")),
            ),
            math_run(" = "),
            math_sup(math_run("e"), math_run("−1")),
            math_fraction(sub("a", "g"), sub("a", "c")),
            math_run(" ≈ 0.368 "),
            math_fraction(sub("a", "g"), sub("a", "c")),
        ],
    )
    add_callout(
        doc,
        "重要修正",
        "缩小约 63% 的是证据相对竞争 token 的 attention odds（相对优势），"
        "不是证据 attention mass 本身必然下降 63%。实际 mass 还取决于其余所有 token 的 softmax 分母。",
    )

    add_heading(doc, "5.2 Attention 决定真实证据写入残差流的“音量”", 2)
    add_body(
        doc,
        "Value 向量携带被检索到的内容，attention weight 决定该内容以多大强度写入当前查询位置。"
        "对一个 head：",
    )
    add_equation(
        doc,
        [
            sub("o", "h"),
            math_run(" = "),
            sub("a", "g,h"),
            math_run(" "),
            sub("v", "g,h"),
            math_run(" + "),
            math_sub(math_run("∑"), math_run("j≠g", plain=True)),
            math_run(" "),
            sub("a", "j,h"),
            math_run(" "),
            sub("v", "j,h"),
        ],
    )
    add_body(
        doc,
        "多个 head 的输出经过 WO 投影，并与原残差及 MLP 更新共同形成下一层表示：",
    )
    add_equation(
        doc,
        [
            sub("r", "ℓ+1"),
            math_run(" = "),
            sub("r", "ℓ"),
            math_run(" + "),
            sub("W", "O"),
            math_run("["),
            sub("o", "1"),
            math_run("; … ; "),
            sub("o", "H"),
            math_run("] + MLP(·)", plain=True),
        ],
    )
    add_body(
        doc,
        "因此，ag,h 降低时，真实 nine 的 Value 信息进入残差流的强度随之减弱；"
        "干扰信息和格式信息在当前表示中的相对占比上升。",
    )

    add_heading(doc, "5.3 最终隐藏状态被 LM Head 读成 nine 或 newline", 2)
    add_body(
        doc,
        "经过所有层后，最终查询状态经 RMSNorm 得到 r̃L。"
        "LM Head 中 nine 与 newline 两行权重的投影差决定最终 margin：",
    )
    add_equation(
        doc,
        [
            sub("Δ", "out"),
            math_run(" = "),
            math_sup(
                [
                    math_run("("),
                    sub("w", "nine"),
                    math_run(" − "),
                    sub("w", "newline"),
                    math_run(")"),
                ],
                math_run("T", plain=True),
            ),
            math_run(" "),
            sub("r̃", "L"),
        ],
    )
    add_table(
        doc,
        ["总长度", "P(nine)", "P(newline)", "输出 margin", "首 token"],
        [
            ["136K", "28.70%", "17.41%", "+0.50", "nine"],
            ["144K", "8.27%", "22.48%", "−1.00", "newline"],
        ],
        [1500, 1800, 2000, 2000, 2060],
        numeric_columns={0, 1, 2, 3},
    )
    add_body(
        doc,
        "从 136K 到 144K，P(nine) 从 28.70% 降至 8.27%，"
        "P(newline) 从 17.41% 升至 22.48%，输出 margin 从 +0.50 变为 −1.00。"
        "当 margin 穿过零点时，newline 即接管首 token。",
    )

    add_heading(doc, "6. 最终判断", 1)
    add_callout(
        doc,
        "因果链",
        "长度与新增内容变化 → Query 状态和 RoPE 相对位置变化 → "
        "真实 nine 的关键 QK 降低 → 证据 attention 相对优势下降 → "
        "nine 的 Value 写入残差流减弱 → 输出 margin 穿过 0 → newline 获胜。",
    )
    add_body(
        doc,
        "关键 QK≈14.11 是当前最有价值的内部预警阈值；"
        "它将失败风险从约 4.8% 提升到约 67.2%，但不能替代输出层的严格判定。"
        "模型真正执行的硬边界始终是：",
    )
    add_equation(
        doc,
        [
            sub("z", "nine"),
            math_run(" − "),
            sub("z", "newline"),
            math_run(" = 0"),
        ],
        after=4,
    )

    add_heading(doc, "附：实验口径", 2)
    add_body(
        doc,
        "序列长度按 Ki token 计：136K=139,264，144K=147,456。"
        "公共 136K 前缀只执行一次 prefill，随后逐 token 更新 KV cache；"
        "每个长度点重新追加固定 23-token 查询并记录 P(nine)、P(newline)、"
        "输出 margin、Gold PPL，以及 29 个关键 head 的加权 QK 与 Query 方向指标。",
        after=0,
    )
    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.core_properties.title = "从 nine 到 newline：证据主导权如何被夺走"
    document.core_properties.subject = "长上下文检索失败边界、关键 QK 软阈值与输出硬边界"
    document.core_properties.author = "Qwen3 long-context retrieval research"
    document.core_properties.keywords = (
        "Qwen3-8B, long context, attention, QK, soft threshold, newline"
    )
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()

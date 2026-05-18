from __future__ import annotations

from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


OUT_DIR = Path("generated_resume")
DOCX_PATH = OUT_DIR / "Yiming_Luo_Resume_EN.docx"
PDF_PATH = OUT_DIR / "Yiming_Luo_Resume_EN.pdf"

INK = RGBColor(17, 24, 39)
MUTED = RGBColor(79, 88, 103)
BLUE = RGBColor(31, 77, 120)
LIGHT_BLUE = RGBColor(232, 238, 245)


resume = {
    "name": "Yiming Luo",
    "headline": "M.S. Student in Computer Science | LLM Systems, KV Cache Compression, Low-Latency C++",
    "contact": "Shanghai, China | 18817438862 | 2781497944@qq.com | WeChat: LittleXi",
    "sections": [
        {
            "title": "Education",
            "items": [
                {
                    "heading": "Fudan University, School of Computer Science",
                    "subheading": "M.S. in Computer Science and Technology",
                    "date": "2025 - 2028 expected",
                    "bullets": [
                        "Research focus: LLM inference systems, long-context memory, KV cache compression, sparse attention, and retrieval-oriented model diagnostics.",
                    ],
                },
                {
                    "heading": "Fudan University, School of Computer Science",
                    "subheading": "B.S. in Computer Science and Technology",
                    "date": "2021 - 2025",
                    "bullets": [
                        "Ranked 8th citywide by raw score in the 2021 science-track Gaokao.",
                    ],
                },
            ],
        },
        {
            "title": "Technical Skills",
            "compact": [
                ("Languages", "C/C++, Python, Bash"),
                ("ML systems", "PyTorch, Hugging Face Transformers, torchrun/DDP, Qwen3, KV cache analysis, sparse attention"),
                ("Systems", "Linux, CMake, shared-library generation, benchmark tooling, Catch2, low-latency C++ optimization"),
                ("Retrieval and tooling", "RAG, embedding retrieval, BM25, VS Code extension development, code-symbol extraction"),
            ],
        },
        {
            "title": "Research Experience",
            "items": [
                {
                    "heading": "Long-Context LLM Systems and KV Cache Compression Research",
                    "subheading": "Independent/Fudan research workspace",
                    "date": "Jan 2026 - May 2026",
                    "bullets": [
                        "Built a modular Qwen3 research workspace that treats KV cache as indexed, compressible, retrievable memory rather than a flat token sequence scanned at every decode step.",
                        "Implemented Qwen3-0.6B/8B profiling pipelines for K-cache norms, K/V value deltas, token-token cosine matrices, SVD/PCA energy, and attention-weighted reconstruction error on DCLM prefixes.",
                        "Measured 28 layers x 8 KV heads over 5,000-token prefixes; found 40/224 heads with off-diagonal K-cache cosine >= 0.9 and 67/224 heads >= 0.8, motivating per-layer/per-head compression policies.",
                        "Ran attention-energy pruning on a 3,000-token DCLM sample: retaining 90% attention energy increased loss by only 0.0142 (PPL 1.0143x), while retaining 95% increased loss by 0.0020 (PPL 1.0020x).",
                        "Developed sparse decode and chunk-routing prototypes, including top-k block selection for layers 3-27, 20-chunk oracle/router experiments, and 8-GPU torchrun/DDP training scripts.",
                        "Prototyped pyramid KV compression with learned weighted K/V summaries, anchor/recent raw-token preservation, RoPE-aware handling, and staged compressor/attention/full-parameter training modes.",
                    ],
                },
            ],
        },
        {
            "title": "Industry Experience",
            "items": [
                {
                    "heading": "Shanghai Boxiong Asset Management Center",
                    "subheading": "C++ Development Engineer",
                    "date": "Jan 2025 - Mar 2025",
                    "bullets": [
                        "Optimized LightGBM decision-tree inference for high-frequency trading with cache-friendly prediction code, JIT-style code generation, and automated CMake/Python generation of shared libraries.",
                        "Reduced single-prediction latency from about 3,200 ns to 803 ns, delivering roughly 4x faster inference in the production prediction path.",
                        "Built benchmark and Catch2 correctness tests, compared against the open-source lleaves implementation, and parallelized shared-library generation workflows.",
                        "Added GPComposer operators and implemented C++ parsing for Shenzhen Stock Exchange binary UDP market-data packets.",
                    ],
                },
                {
                    "heading": "Shanghai Quantitative Hedge Fund",
                    "subheading": "Low-Latency Trading System and TTS Test Development",
                    "date": "Feb 2024 - May 2024",
                    "bullets": [
                        "Developed low-latency trading-system components and transaction-test tooling for financial trading workflows.",
                        "Worked on performance-sensitive C++ code paths and automated tests for market-data and trading-system integration.",
                    ],
                },
            ],
        },
        {
            "title": "Selected Projects",
            "items": [
                {
                    "heading": "RAG Knowledge Base and VS Code Code-Intelligence Plugin",
                    "subheading": "LLM tooling project",
                    "date": "Mar 2025 - Apr 2025",
                    "bullets": [
                        "Developed a VS Code extension that extracts source-code symbols and function bodies for downstream retrieval and code understanding.",
                        "Combined embedding retrieval and BM25 to evaluate similarity across documentation and code repositories; built domain knowledge-base workflows supporting build and query modes.",
                    ],
                },
                {
                    "heading": "AI Agent for Memorial Conversation and Voice Chat",
                    "subheading": "AI Engineer, Shiguang",
                    "date": "Mar 2026 - May 2026",
                    "bullets": [
                        "Built an AI-agent prototype that models a deceased family member's biography, memories, and conversational style for text and voice-based interaction.",
                        "Integrated structured life-history data with dialogue behavior to support personalized, memory-grounded conversations.",
                    ],
                },
                {
                    "heading": "Autonomous Driving Safety Testing Platform",
                    "subheading": "Full-stack and simulation development",
                    "date": "Dec 2022 - May 2023",
                    "bullets": [
                        "Extended the SVL autonomous-driving simulator and built Unity3D simulation scenes for automated safety testing.",
                        "Integrated front-end/back-end platform components with fuzzing-based test generation to support automated driving-safety evaluation pipelines.",
                    ],
                },
                {
                    "heading": "RISC-V Operating System Lab",
                    "subheading": "MIT 6.S081-style systems project",
                    "date": "Sep 2023 - Jan 2024",
                    "bullets": [
                        "Implemented core OS mechanisms including process creation, scheduling, termination, inter-process communication, copy-on-write memory, and file-system features such as symbolic links.",
                    ],
                },
            ],
        },
        {
            "title": "Publication",
            "items": [
                {
                    "heading": "Limitations and Future Directions of \"Data Shapley: Equitable Valuation of Data for Machine Learning\"",
                    "subheading": "Well Testing Journal",
                    "date": "Sep 2024",
                    "bullets": [
                        "Paper URL: https://welltestingjournal.com/index.php/WT/article/view/Limitations_and_future_directions_of_Data_Shapley_Equitable_Valu",
                    ],
                }
            ],
        },
        {
            "title": "Teaching and Competitive Programming",
            "items": [
                {
                    "heading": "Algorithm Instructor and Competitive Programming Coach",
                    "subheading": "Online/offline one-on-one teaching",
                    "date": "2021 - Present",
                    "bullets": [
                        "Coached 40+ students across Mainland China, Hong Kong, the United States, Canada, and Australia for CSP-J/S, NOIP, USACO, and competitive-programming training.",
                        "Built an online algorithm-learning presence with 6,000+ followers and published technical lectures, including a tree-chain decomposition tutorial.",
                    ],
                },
                {
                    "heading": "Competitive Programming Profile",
                    "subheading": "Algorithms and contests",
                    "date": "",
                    "bullets": [
                        "Codeforces International Master, rating 2300+ and top 0.5% globally; LeetCode rating 2600+ with global top-1,000 ranking.",
                        "Accumulated 300,000+ lines of C++/Python implementation work across contests, systems projects, and research prototypes.",
                    ],
                },
            ],
        },
        {
            "title": "Honors",
            "bullets": [
                "Gold Medal, 49th ACM-ICPC Asia Regional Contest, Nov 2024.",
                "Gold Medal, 10th China Collegiate Programming Contest Regional Contest, Oct 2024.",
                "Gold Medal, 2024 CCPC Shanghai Municipal Contest, Jun 2024.",
                "Silver Medal, 48th ACM-ICPC Asia East Continent Final, Nov 2023.",
                "Silver Medal, 48th ACM-ICPC Asia Regional Contest, Nanjing, Dec 2023.",
                "Silver Medal, ACM-ICPC Asia Regional Contest, Hangzhou, Oct 2024.",
            ],
        },
    ],
}


def set_docx_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_border_bottom(paragraph, color="D7DBE2", size="6", space="3"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_docx_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    set_docx_font(run, size=10.5, bold=True, color=BLUE)
    set_paragraph_border_bottom(p)


def add_docx_entry(doc, heading, subheading, date, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(heading)
    set_docx_font(run, size=9.5, bold=True, color=INK)
    if subheading:
        run = p.add_run(f" | {subheading}")
        set_docx_font(run, size=9.2, color=INK)
    if date:
        run = p.add_run(f" | {date}")
        set_docx_font(run, size=9.0, italic=True, color=MUTED)
    for bullet in bullets:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.left_indent = Inches(0.20)
        bp.paragraph_format.first_line_indent = Inches(-0.11)
        bp.paragraph_format.space_after = Pt(1.3)
        bp.paragraph_format.line_spacing = 1.0
        br = bp.add_run(bullet)
        set_docx_font(br, size=8.75, color=INK)


def add_docx_compact_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.35)
    table.columns[1].width = Inches(5.95)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(1.35)
        cells[1].width = Inches(5.95)
        cells[0].text = ""
        cells[1].text = ""
        p0 = cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        set_docx_font(r0, size=8.7, bold=True, color=BLUE)
        p1 = cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(value)
        set_docx_font(r1, size=8.7, color=INK)
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = OxmlElement(f"w:{edge}")
                tag.set(qn("w:val"), "nil")
                tc_borders.append(tag)
            tc_pr.append(tc_borders)


def build_docx():
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.50)
    section.bottom_margin = Inches(0.50)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(resume["name"])
    set_docx_font(run, size=20, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(resume["headline"])
    set_docx_font(run, size=9.5, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(resume["contact"])
    set_docx_font(run, size=8.5, color=MUTED)

    for section_data in resume["sections"]:
        add_docx_section_heading(doc, section_data["title"])
        if "compact" in section_data:
            add_docx_compact_table(doc, section_data["compact"])
        if "items" in section_data:
            for item in section_data["items"]:
                add_docx_entry(
                    doc,
                    item["heading"],
                    item.get("subheading", ""),
                    item.get("date", ""),
                    item.get("bullets", []),
                )
        if "bullets" in section_data:
            for bullet in section_data["bullets"]:
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.left_indent = Inches(0.20)
                bp.paragraph_format.first_line_indent = Inches(-0.11)
                bp.paragraph_format.space_after = Pt(1.3)
                br = bp.add_run(bullet)
                set_docx_font(br, size=8.75, color=INK)

    for sec in doc.sections:
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer.add_run("Yiming Luo Resume")
        set_docx_font(footer_run, size=7.5, color=MUTED)

    doc.save(DOCX_PATH)


def pdf_styles():
    base = getSampleStyleSheet()
    return {
        "name": ParagraphStyle(
            "Name",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=19,
            leading=22,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#111827"),
            spaceAfter=2,
        ),
        "headline": ParagraphStyle(
            "Headline",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=9.3,
            leading=11,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#1F4D78"),
            spaceAfter=2,
        ),
        "contact": ParagraphStyle(
            "Contact",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.2,
            leading=10,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#4F5867"),
            spaceAfter=6,
        ),
        "section": ParagraphStyle(
            "Section",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10.2,
            leading=12,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=5,
            spaceAfter=2,
            keepWithNext=True,
        ),
        "entry_left": ParagraphStyle(
            "EntryLeft",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=9.2,
            leading=10.6,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#111827"),
            spaceAfter=0,
        ),
        "entry_date": ParagraphStyle(
            "EntryDate",
            parent=base["Normal"],
            fontName="Helvetica-Oblique",
            fontSize=8.2,
            leading=10,
            alignment=TA_RIGHT,
            textColor=colors.HexColor("#4F5867"),
            spaceAfter=0,
        ),
        "bullet": ParagraphStyle(
            "Bullet",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.15,
            leading=9.35,
            leftIndent=8,
            firstLineIndent=0,
            textColor=colors.HexColor("#111827"),
            spaceAfter=1.4,
        ),
        "label": ParagraphStyle(
            "Label",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=8.25,
            leading=9.8,
            textColor=colors.HexColor("#1F4D78"),
        ),
        "compact": ParagraphStyle(
            "Compact",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.25,
            leading=9.8,
            textColor=colors.HexColor("#111827"),
        ),
    }


def add_pdf_section(story, styles, title):
    story.append(Paragraph(title.upper(), styles["section"]))
    story.append(HRFlowable(width="100%", thickness=0.45, color=colors.HexColor("#D7DBE2"), spaceBefore=0, spaceAfter=2))


def add_pdf_entry(story, styles, heading, subheading, date, bullets):
    left = f"<b>{escape(heading)}</b>"
    if subheading:
        left += f" | {escape(subheading)}"
    table = Table(
        [[Paragraph(left, styles["entry_left"]), Paragraph(escape(date), styles["entry_date"])]],
        colWidths=[5.25 * inch, 1.25 * inch],
        hAlign="LEFT",
        style=TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        ),
    )
    story.append(table)
    list_items = [
        ListItem(Paragraph(escape(bullet), styles["bullet"]), leftIndent=4, bulletColor=colors.HexColor("#1F4D78"))
        for bullet in bullets
    ]
    story.append(ListFlowable(list_items, bulletType="bullet", start="circle", leftIndent=12, bulletFontSize=4.8))
    story.append(Spacer(1, 1.4))


def build_pdf():
    OUT_DIR.mkdir(exist_ok=True)
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=LETTER,
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.46 * inch,
        bottomMargin=0.42 * inch,
        title="Yiming Luo Resume",
        author="Yiming Luo",
    )
    story = [
        Paragraph(resume["name"], styles["name"]),
        Paragraph(escape(resume["headline"]), styles["headline"]),
        Paragraph(escape(resume["contact"]), styles["contact"]),
    ]
    for section_data in resume["sections"]:
        add_pdf_section(story, styles, section_data["title"])
        if "compact" in section_data:
            rows = [
                [
                    Paragraph(escape(label), styles["label"]),
                    Paragraph(escape(value), styles["compact"]),
                ]
                for label, value in section_data["compact"]
            ]
            table = Table(rows, colWidths=[1.25 * inch, 5.25 * inch], hAlign="LEFT")
            table.setStyle(
                TableStyle(
                    [
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 0),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                        ("TOPPADDING", (0, 0), (-1, -1), 0),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 2))
        if "items" in section_data:
            for item in section_data["items"]:
                add_pdf_entry(
                    story,
                    styles,
                    item["heading"],
                    item.get("subheading", ""),
                    item.get("date", ""),
                    item.get("bullets", []),
                )
        if "bullets" in section_data:
            list_items = [
                ListItem(Paragraph(escape(bullet), styles["bullet"]), leftIndent=4, bulletColor=colors.HexColor("#1F4D78"))
                for bullet in section_data["bullets"]
            ]
            story.append(ListFlowable(list_items, bulletType="bullet", start="circle", leftIndent=12, bulletFontSize=4.8))

    def footer(canvas, doc_obj):
        canvas.saveState()
        canvas.setFont("Helvetica", 7.3)
        canvas.setFillColor(colors.HexColor("#6B7280"))
        canvas.drawRightString(LETTER[0] - 0.55 * inch, 0.25 * inch, f"Yiming Luo Resume | Page {doc_obj.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOCX_PATH.resolve())
    print(PDF_PATH.resolve())
